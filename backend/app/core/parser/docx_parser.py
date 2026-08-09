"""DOCX 解析器。"""
from docx import Document as DocxDocument

from app.core.parser.models import ParsedElement


async def parse(path: str) -> list[ParsedElement]:
    """解析 docx：段落（Heading N → heading + level；其余 text）+ 表格（HTML）。"""
    elements: list[ParsedElement] = []
    d = DocxDocument(path)
    for para in d.paragraphs:
        txt = (para.text or "").strip()
        if not txt:
            continue
        style = (para.style.name or "").lower()
        if style.startswith("heading"):
            try:
                lvl = int(style.split()[-1])
            except ValueError:
                lvl = 1
            elements.append(ParsedElement("heading", txt, 1, level=lvl))
        elif style == "title":
            elements.append(ParsedElement("heading", txt, 1, level=1))
        else:
            elements.append(ParsedElement("text", txt, 1))
    for table in d.tables:
        rows = [[c.text for c in row.cells] for row in table.rows]
        body = "".join("<tr>" + "".join(f"<td>{c}</td>" for c in r) + "</tr>" for r in rows)
        elements.append(ParsedElement("table", f"<table>{body}</table>", 1))
    return elements
